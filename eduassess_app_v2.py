
import os
from io import BytesIO

import streamlit as st
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
from openai import OpenAI


load_dotenv()

st.set_page_config(
    page_title="EduAssess AI Agent V2",
    page_icon="📝",
    layout="wide",
)

st.title("📝 EduAssess AI Agent V2")

st.markdown(
    """
### Intelligent Assessment Generator

Create professional classroom assessments using **Generative AI** and **Agentic AI**.

- Multiple question types
- Automatic answer key
- Bloom’s Taxonomy classification
- Ready-to-use classroom assessments
"""
)

col1, col2, col3 = st.columns(3)

with col1:
    st.metric("Question Types", "4")

with col2:
    st.metric("Maximum Questions", "30")

with col3:
    st.metric("AI Mode", "Demo + API")

st.success(
    "🟢 System Ready — Demo Mode is active. "
    "OpenAI API support is available when an API key is configured."
)

st.write(
    """
This application helps teachers generate quizzes and exams using
Generative and Agentic AI.

The agent receives assessment requirements and creates questions,
answer keys, and Bloom's Taxonomy levels.
"""
)

with st.sidebar:
    st.header("Assessment Information")

    grade = st.text_input("Grade Level", "Grade 8")
    subject = st.text_input("Subject", "ICT")
    topic = st.text_input("Topic", "Cybersecurity")

    number_questions = st.slider(
        "Number of Questions",
        min_value=5,
        max_value=30,
        value=10,
    )

    difficulty = st.selectbox(
        "Difficulty Level",
        ["Easy", "Medium", "Hard", "Mixed"],
    )

    question_types = st.multiselect(
        "Question Types",
        [
            "Multiple Choice",
            "True/False",
            "Short Answer",
            "Matching",
        ],
        default=[
            "Multiple Choice",
            "True/False",
            "Short Answer",
        ],
    )

    generate = st.button(
        "Generate Assessment",
        type="primary",
        use_container_width=True,
    )
     
st.divider()

estimated_time = number_questions * 2

st.subheader("📋 Assessment Summary")

col1, col2 = st.columns(2)

with col1:
    st.write(f"**Grade:** {grade}")
    st.write(f"**Subject:** {subject}")
    st.write(f"**Topic:** {topic}")

with col2:
    st.write(f"**Difficulty:** {difficulty}")
    st.write(f"**Questions:** {number_questions}")
    st.write(f"**Estimated Time:** {estimated_time} minutes") 
view_mode = st.radio(
    "👁️ View Mode",
    ["👨‍🏫 Teacher View", "👨‍🎓 Student View"],
    horizontal=True,
)

def demo_assessment(
    grade,
    subject,
    topic,
    number_questions,
    difficulty,
    question_types,
):
    questions = []

    for i in range(1, number_questions + 1):
        if i % 3 == 1:
            questions.append(
                f"""
**Q{i}. Multiple Choice:** Which statement best describes {topic}?

A. It is unrelated to digital safety.  
B. It helps protect systems, data, and users.  
C. It only applies to hardware.  
D. It is used only in gaming.

**Answer:** B  
**Bloom's Level:** Understanding  
**Difficulty:** {difficulty}
"""
            )

        elif i % 3 == 2:
            questions.append(
                f"""
**Q{i}. True/False:** {topic} is important for protecting personal
and organizational information.

**Answer:** True  
**Bloom's Level:** Remembering  
**Difficulty:** {difficulty}
"""
            )

        else:
            questions.append(
                f"""
**Q{i}. Short Answer:** Explain one real-life example of {topic}.

**Sample Answer:** A real-life example is using strong passwords and
two-factor authentication to protect online accounts.

**Bloom's Level:** Applying  
**Difficulty:** {difficulty}
"""
            )

    selected_types = (
        ", ".join(question_types)
        if question_types
        else "No question type selected"
    )

    return f"""
# Assessment: {topic}

## Assessment Information

- **Grade:** {grade}
- **Subject:** {subject}
- **Topic:** {topic}
- **Number of Questions:** {number_questions}
- **Difficulty:** {difficulty}
- **Question Types:** {selected_types}

## Student Instructions

1. Read every question carefully.
2. Answer all questions.
3. Select only one answer for each multiple-choice question.
4. Review your answers before submission.

## Questions

{''.join(questions)}

## Assessment Criteria

- Accuracy of answers
- Understanding of key concepts
- Application to real-life situations
- Critical thinking
- Clear communication

## Agentic AI Role

This AI agent receives teacher-defined goals, analyzes assessment
requirements, organizes the assessment structure, and generates
questions, answers, and learning-level classifications.
"""
def create_student_view(content):
    hidden_lines = []

    for line in content.splitlines():
        if line.startswith("**Answer:**"):
            continue
        if line.startswith("**Sample Answer:**"):
            continue
        if line.startswith("**Bloom's Level:**"):
            continue
        if line.startswith("**Difficulty:**"):
            continue

        hidden_lines.append(line)

    return "\n".join(hidden_lines)

def create_pdf(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    for line in content.splitlines():
        clean_line = (
            line.replace("**", "")
            .replace("#", "")
            .replace("📝", "")
            .replace("👨‍🏫", "")
            .replace("👨‍🎓", "")
            .replace("—", "-")
            .replace("’", "'")
            .replace("“", '"')
            .replace("”", '"')
        )

        # Standart PDF şriftinin dəstəkləmədiyi simvolları təmizləyir
        clean_line = clean_line.encode(
            "latin-1",
            errors="replace"
        ).decode("latin-1")

        pdf.multi_cell(
            w=0,
            h=7,
            text=clean_line if clean_line else " ",
            new_x="LMARGIN",
            new_y="NEXT",
        )

    return bytes(pdf.output())
def create_docx(content):
    document = Document()

    for line in content.splitlines():
        clean_line = (
            line.replace("**", "")
            .replace("#", "")
            .replace("📝", "")
            .replace("👨‍🏫", "")
            .replace("👨‍🎓", "")
        )

        if clean_line.strip():
            document.add_paragraph(clean_line)
        else:
            document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()

def generate_with_openai(prompt):
    api_key = os.getenv("OPENAI_API_KEY")

    if not api_key:
        return None

    try:
        client = OpenAI(api_key=api_key)

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert educational assessment "
                        "designer and Agentic AI assistant."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.5,
        )

        return response.choices[0].message.content

    except Exception as error:
        st.warning(
            "OpenAI API could not be used. "
            f"Demo mode activated. Error: {error}"
        )
        return None


if generate:
    if not question_types:
        st.warning("Please select at least one question type.")

    else:
        with st.spinner("🤖 AI is analyzing your inputs and generating the assessment..."):
            prompt = f"""
Create a professional educational assessment.

Grade: {grade}
Subject: {subject}
Topic: {topic}
Number of Questions: {number_questions}
Difficulty: {difficulty}
Question Types: {", ".join(question_types)}

Include:

1. Assessment title
2. Instructions for students
3. Questions
4. Answer key
5. Bloom's Taxonomy level for each question
6. Difficulty level for each question
7. Marking criteria
8. Explanation of how this represents Agentic AI

Use clear, age-appropriate, and academically correct language.
"""

            result = generate_with_openai(prompt)

            if result is None:
                result = demo_assessment(
                    grade,
                    subject,
                    topic,
                    number_questions,
                    difficulty,
                    question_types,
                )

        st.success("Assessment generated successfully.")

        if view_mode == "👨‍🎓 Student View":
                display_result = create_student_view(result)
        else:
                display_result = result

        st.markdown(display_result)

        st.download_button(
            label="Download Assessment as TXT",
            data=display_result,
            file_name="assessment.txt",
            mime="text/plain",
            use_container_width=True,
        )

        pdf_data = create_pdf(display_result)

        st.download_button(
            label="📄 Download Assessment as PDF",
            data=pdf_data,
            file_name="assessment.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
        docx_data = create_docx(display_result)

        st.download_button(
            label="📘 Download Assessment as Word",
            data=docx_data,
            file_name="assessment.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
else:
    st.info(
        "Enter assessment information in the sidebar and click "
        "'Generate Assessment'."
    )
